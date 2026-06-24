from pathlib import Path

from app.parsers.base import ParseResult, ParserAdapter
from app.parsers.docx import DocxParser
from app.parsers.html import HtmlParser
from app.parsers.markdown import MarkdownParser
from app.parsers.pdf import PdfParser
from app.parsers.text import TextParser


def test_markdown_parser_extracts_title_sections_and_links(tmp_path) -> None:
    path = tmp_path / "rag.md"
    path.write_text(
        "# RAG 笔记\n\n"
        "检索增强生成。\n\n"
        "## Agent\n\n"
        "参考 [OpenAI](https://openai.com)。\n",
        encoding="utf-8",
    )

    parser = MarkdownParser()
    result = parser.parse(path)

    assert isinstance(parser, ParserAdapter)
    assert isinstance(result, ParseResult)
    assert result.title == "RAG 笔记"
    assert "检索增强生成" in result.text
    assert [section.heading for section in result.sections] == ["RAG 笔记", "Agent"]
    assert result.links == ["https://openai.com"]
    assert result.metadata["source_format"] == "markdown"
    assert parser.can_parse(path, "text/markdown", {})


def test_text_parser_reads_plain_text_and_warns_for_empty_file(tmp_path) -> None:
    path = tmp_path / "plain.txt"
    path.write_text("第一行\n第二行", encoding="utf-8")
    empty_path = tmp_path / "empty.txt"
    empty_path.write_text("", encoding="utf-8")

    parser = TextParser()
    result = parser.parse(path)
    empty_result = parser.parse(empty_path)

    assert result.title == "plain"
    assert result.text == "第一行\n第二行"
    assert result.sections == []
    assert empty_result.text == ""
    assert "empty_file" in empty_result.warnings


def test_text_parser_returns_warning_for_encoding_errors(tmp_path) -> None:
    path = tmp_path / "broken.txt"
    path.write_bytes(b"\xff\xfe\x00")

    result = TextParser().parse(path)

    assert result.text == ""
    assert "decode_error" in result.warnings
    assert result.metadata["source_format"] == "text"


def test_pdf_parser_extracts_text_and_page_map(tmp_path) -> None:
    path = tmp_path / "sample.pdf"
    _write_pdf(path, ["page one knowledge", "page two knowledge"])

    result = PdfParser().parse(path)

    assert result.title == "sample"
    assert "page one knowledge" in result.text
    assert "page two knowledge" in result.text
    assert [page.page_number for page in result.page_map] == [1, 2]
    assert result.metadata["source_format"] == "pdf"


def test_docx_parser_extracts_heading_and_paragraphs(tmp_path) -> None:
    path = tmp_path / "note.docx"
    _write_docx(path)

    result = DocxParser().parse(path)

    assert result.title == "知识库标题"
    assert "正文段落" in result.text
    assert [section.heading for section in result.sections] == ["知识库标题"]
    assert result.metadata["source_format"] == "docx"


def test_html_parser_extracts_title_body_and_links(tmp_path) -> None:
    path = tmp_path / "page.html"
    path.write_text(
        "<html><head><title>网页标题</title><script>ignored()</script></head>"
        "<body><article><h1>文章标题</h1><p>正文内容</p>"
        "<a href=\"https://example.com/a\">链接</a></article></body></html>",
        encoding="utf-8",
    )

    result = HtmlParser().parse(path)

    assert result.title == "网页标题"
    assert "文章标题" in result.text
    assert "正文内容" in result.text
    assert "ignored" not in result.text
    assert result.links == ["https://example.com/a"]
    assert result.metadata["source_format"] == "html"


def _write_pdf(path: Path, page_texts) -> None:
    import fitz

    document = fitz.open()
    for text in page_texts:
        page = document.new_page()
        page.insert_text((72, 72), text)
    document.save(path)
    document.close()


def _write_docx(path: Path) -> None:
    from docx import Document

    document = Document()
    document.add_heading("知识库标题", level=1)
    document.add_paragraph("正文段落")
    document.save(path)
